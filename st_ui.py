import streamlit as st
import os
import tempfile
import logging
from datetime import datetime
from docx import Document
from io import BytesIO
from openai import OpenAI

# Import all functions from f3.py
from f3 import (
    iter_block_items, extract_articles, merge_gazette, apply_amendment_text,
    extract_amending_ref, add_explanatory_table, deep_copy_paragraph,
    deep_copy_table, ARTICLE_RE, CHANGE_RE, hex_to_rgb, set_alignment,
    RGBColor, Pt, SequenceMatcher
)

# Initialize OpenAI client with secrets
try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except KeyError:
    st.error("⚠️ OpenAI API key not found in secrets.toml. Please add OPENAI_API_KEY to your secrets configuration.")
    client = None

# Setup logging for Streamlit
if 'log_messages' not in st.session_state:
    st.session_state.log_messages = []

def add_log_message(message):
    timestamp = datetime.now().strftime('%H:%M:%S')
    st.session_state.log_messages.append(f"{timestamp} - {message}")

def process_part_a(orig_file, amend_file):
    """Process Part A: Generate New.docx"""
    if not client:
        add_log_message("Error: OpenAI API key not configured")
        return None
        
    add_log_message("Processing Part A...")
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_orig:
            tmp_orig.write(orig_file.read())
            orig_path = tmp_orig.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_amend:
            tmp_amend.write(amend_file.read())
            amend_path = tmp_amend.name
        
        # Load documents
        orig = Document(orig_path)
        updated_doc = Document()
        
        # Copy original document structure
        for block in iter_block_items(orig):
            if hasattr(block, 'text'):  # Paragraph
                p = updated_doc.add_paragraph()
                deep_copy_paragraph(block, p)
            else:  # Table
                deep_copy_table(block, updated_doc)
        
        # Merge gazette information
        merge_gazette(orig, Document(amend_path), updated_doc)
        
        # Extract articles and apply amendments
        articles = extract_articles(updated_doc)
        amend_doc = Document(amend_path)
        amendments = [p.text.strip() for p in amend_doc.paragraphs if "prestaju da važe" in p.text]
        
        for inst in amendments:
            matches = CHANGE_RE.findall(inst)
            for article_num, stav_num in matches:
                aid = f"Član {article_num}"
                if aid in articles:
                    s, e = articles[aid]
                    old_lines = '\n'.join(updated_doc.paragraphs[k].text for k in range(s, e))
                    new_lines = apply_amendment_text(old_lines, inst)
                    
                    # Remove old paragraphs
                    for k in range(e - 1, s - 1, -1):
                        updated_doc.paragraphs[k]._element.getparent().remove(updated_doc.paragraphs[k]._element)
                    
                    # Insert new paragraphs
                    for i, line in enumerate(new_lines):
                        new_p = updated_doc.add_paragraph(line)
                        original_format_idx = s + i
                        if original_format_idx < len(orig.paragraphs):
                            deep_copy_paragraph(orig.paragraphs[original_format_idx], new_p)
                        
                        # Format article titles
                        if ARTICLE_RE.match(line):
                            set_alignment(new_p, 'center')
                            for run in new_p.runs:
                                run.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
        
        # Clean up temporary files
        os.unlink(orig_path)
        os.unlink(amend_path)
        
        add_log_message("Part A completed successfully")
        return updated_doc
        
    except Exception as e:
        add_log_message(f"Error in Part A: {str(e)}")
        return None

def process_part_b(orig_file, new_file, amend_file):
    """Process Part B: Generate Colored Diff.docx"""
    add_log_message("Processing Part B...")
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_orig:
            tmp_orig.write(orig_file.read())
            orig_path = tmp_orig.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_new:
            tmp_new.write(new_file.read())
            new_path = tmp_new.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_amend:
            tmp_amend.write(amend_file.read())
            amend_path = tmp_amend.name
        
        # Load documents
        orig = Document(orig_path)
        new_d = Document(new_path)
        gov = Document(amend_path)
        diff_doc = Document()
        
        # Add explanatory table
        add_explanatory_table(diff_doc)
        
        # Prepare texts for diff
        orig_blocks = list(iter_block_items(orig))
        new_blocks = list(iter_block_items(new_d))
        
        orig_texts = []
        for b in orig_blocks:
            if hasattr(b, 'text'):
                orig_texts.append(b.text.strip())
            else:
                orig_texts.append('TABLE')
        
        new_texts = []
        for b in new_blocks:
            if hasattr(b, 'text'):
                text = b.text.strip().replace('*', '')
                new_texts.append(text)
            else:
                new_texts.append('TABLE')
        
        # Generate diff
        matcher = SequenceMatcher(None, orig_texts, new_texts)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                for k in range(i1, i2):
                    block = new_blocks[j1 + (k - i1)]
                    if hasattr(block, 'text'):
                        p = diff_doc.add_paragraph()
                        deep_copy_paragraph(block, p)
                    else:
                        deep_copy_table(block, diff_doc)
            elif tag == 'delete':
                for k in range(i1, i2):
                    block = orig_blocks[k]
                    if hasattr(block, 'text'):
                        p = diff_doc.add_paragraph('[' + block.text + ']')
                        deep_copy_paragraph(block, p)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        deep_copy_table(block, diff_doc, color=RGBColor(255, 0, 0))
            elif tag == 'insert':
                for k in range(j1, j2):
                    block = new_blocks[k]
                    if hasattr(block, 'text'):
                        p = diff_doc.add_paragraph('[' + block.text + ']')
                        deep_copy_paragraph(block, p)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0, 204, 51)
                    else:
                        deep_copy_table(block, diff_doc, color=RGBColor(0, 204, 51))
            elif tag == 'replace':
                for k in range(i1, i2):
                    block = orig_blocks[k]
                    if hasattr(block, 'text'):
                        p = diff_doc.add_paragraph('[' + block.text + ']')
                        deep_copy_paragraph(block, p)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        deep_copy_table(block, diff_doc, color=RGBColor(255, 0, 0))
                for k in range(j1, j2):
                    block = new_blocks[k]
                    if hasattr(block, 'text'):
                        p = diff_doc.add_paragraph('[' + block.text + ']')
                        deep_copy_paragraph(block, p)
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0, 204, 51)
                    else:
                        deep_copy_table(block, diff_doc, color=RGBColor(0, 204, 51))
        
        # Insert dynamic green reference
        diff_blocks = list(iter_block_items(diff_doc))
        insert_positions = []
        ref = extract_amending_ref(gov)
        
        for idx, block in enumerate(diff_blocks):
            if hasattr(block, 'text') and ARTICLE_RE.match(block.text.strip()) and '*' in block.text:
                insert_positions.append(idx)
        
        for pos in reversed(insert_positions):
            ref_p = diff_doc.add_paragraph()
            set_alignment(ref_p, 'center')
            run = ref_p.add_run(ref)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 204, 51)
            diff_doc._body._element.insert(pos, ref_p._element)
        
        # Clean up temporary files
        os.unlink(orig_path)
        os.unlink(new_path)
        os.unlink(amend_path)
        
        add_log_message("Part B completed successfully")
        return diff_doc
        
    except Exception as e:
        add_log_message(f"Error in Part B: {str(e)}")
        return None

def doc_to_bytes(doc):
    """Convert Document to bytes for download"""
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def main():
    st.set_page_config(
        page_title="Legal Document Processor",
        page_icon="⚖️",
        layout="wide"
    )
    
    st.title("⚖️ Legal Document Processor")
    st.markdown("---")
    
    # Initialize session state
    if 'updated_doc' not in st.session_state:
        st.session_state.updated_doc = None
    if 'diff_doc' not in st.session_state:
        st.session_state.diff_doc = None
    
    # File upload section
    st.header("📁 Upload Documents")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.subheader("Original Law")
        orig_file = st.file_uploader(
            "Upload original law document (.docx)",
            type=['docx'],
            key='orig_file'
        )
    
    with col2:
        st.subheader("Government Changes")
        amend_file = st.file_uploader(
            "Upload government changes document (.docx)",
            type=['docx'],
            key='amend_file'
        )
    
    with col3:
        st.subheader("New Law (for diff)")
        new_file = st.file_uploader(
            "Upload new law document (.docx)",
            type=['docx'],
            key='new_file'
        )
    
    st.markdown("---")
    
    # Processing section
    st.header("🔄 Processing")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button(
            "🔧 Process Part A: Generate New.docx",
            disabled=not (orig_file and amend_file),
            use_container_width=True
        ):
            with st.spinner("Processing Part A..."):
                st.session_state.updated_doc = process_part_a(orig_file, amend_file)
            
            if st.session_state.updated_doc:
                st.success("✅ New.docx generated successfully!")
            else:
                st.error("❌ Failed to generate new document")
    
    with col2:
        if st.button(
            "🎨 Process Part B: Generate Colored Diff.docx",
            disabled=not (orig_file and new_file and amend_file),
            use_container_width=True
        ):
            with st.spinner("Processing Part B..."):
                st.session_state.diff_doc = process_part_b(orig_file, new_file, amend_file)
            
            if st.session_state.diff_doc:
                st.success("✅ Colored diff.docx generated successfully!")
            else:
                st.error("❌ Failed to generate diff document")
    
    st.markdown("---")
    
    # Download section
    st.header("💾 Download Results")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.session_state.updated_doc:
            doc_bytes = doc_to_bytes(st.session_state.updated_doc)
            st.download_button(
                label="📥 Download New.docx",
                data=doc_bytes,
                file_name=f"new_law_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("📥 Download New.docx", disabled=True, use_container_width=True)
    
    with col2:
        if st.session_state.diff_doc:
            doc_bytes = doc_to_bytes(st.session_state.diff_doc)
            st.download_button(
                label="📥 Download Colored Diff.docx",
                data=doc_bytes,
                file_name=f"colored_diff_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("📥 Download Colored Diff.docx", disabled=True, use_container_width=True)
    
    st.markdown("---")
    
    # Log section
    st.header("📋 Processing Log")
    
    if st.session_state.log_messages:
        log_container = st.container()
        with log_container:
            for message in st.session_state.log_messages[-10:]:  # Show last 10 messages
                st.text(message)
    else:
        st.info("No log messages yet. Start processing to see updates.")
    
    if st.button("🗑️ Clear Log"):
        st.session_state.log_messages = []
        st.rerun()
    
    # Instructions
    with st.expander("ℹ️ Instructions"):
        st.markdown("""
        ### How to use:
        
        1. **Upload Documents**: Upload the required .docx files:
           - **Original Law**: The base law document
           - **Government Changes**: Document containing amendments
           - **New Law**: Updated law document (needed for diff generation)
        
        2. **Process Part A**: Generate a new law document with amendments applied
        
        3. **Process Part B**: Generate a colored diff showing changes between original and new versions
        
        4. **Download**: Save the processed documents to your computer
        
        ### Color Coding in Diff:
        - **Green**: New content added
        - **Red**: Content removed
        - **Black**: Unchanged content
        """)

if __name__ == "__main__":
    main()
