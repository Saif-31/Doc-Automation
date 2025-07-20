import os
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from openai import OpenAI
from difflib import SequenceMatcher

# Try to import streamlit for cloud deployment, fall back to dotenv for local
try:
    import streamlit as st
    # Running in Streamlit environment
    try:
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    except (KeyError, AttributeError):
        # Fallback to environment variable if secrets not available
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except ImportError:
    # Running locally, use dotenv
    from dotenv import load_dotenv
    load_dotenv()
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Setup logging
logging.basicConfig(filename=f"processor_{datetime.now().strftime('%Y%m%d')}.log", level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Configuration
DEFAULT_DIR = r"./docs/samples"
os.makedirs(DEFAULT_DIR, exist_ok=True)

ARTICLE_RE = re.compile(r"^Član\s+(\d+[a-zA-Z]*)", re.IGNORECASE)
GAZETTE_RE = re.compile(r'\"Sl\. glasnik RS\", br\. (.+?)\)')
CHANGE_RE = re.compile(r"člana (\d+)\. stav (\d+)\. Zakona o računovodstvu")

# Utility Functions
def iter_block_items(doc):
    for child in doc._body._element:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def get_font_color(run):
    c = run.font.color
    if c is not None and c.rgb is not None:
        return f"#{c.rgb[0]:02X}{c.rgb[1]:02X}{c.rgb[2]:02X}"
    return None

def get_run_shading(run):
    r_pr = run._element.rPr
    if r_pr is not None:
        shd = r_pr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            return f"#{fill}" if fill else None
    return None

def get_paragraph_shading(paragraph):
    pPr = paragraph._element.pPr
    if pPr is not None:
        shd = pPr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            return f"#{fill}" if fill else None
    return None

def get_cell_shading(cell):
    tc_pr = cell._tc.tcPr
    if tc_pr is not None:
        shd = tc_pr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            return f"#{fill}" if fill else None
    return None

def get_paragraph_indents(paragraph):
    pf = paragraph.paragraph_format
    d = {}
    if pf.left_indent is not None: d['left'] = pf.left_indent.pt
    if pf.right_indent is not None: d['right'] = pf.right_indent.pt
    if pf.first_line_indent is not None:
        if pf.first_line_indent.pt < 0:
            d['hanging'] = abs(pf.first_line_indent.pt)
        else:
            d['first_line'] = pf.first_line_indent.pt
    return d

def get_paragraph_spacing(paragraph):
    pf = paragraph.paragraph_format
    d = {}
    if pf.space_before is not None: d['space_before'] = pf.space_before.pt
    if pf.space_after is not None: d['space_after'] = pf.space_after.pt
    if pf.line_spacing is not None: d['line_spacing'] = pf.line_spacing
    return d

def set_run_shading(run, shading_hex):
    if not shading_hex:
        return
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shading_hex.lstrip("#"))
    rPr.append(shd)

def set_paragraph_shading(paragraph, shading_hex):
    if not shading_hex:
        return
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shading_hex.lstrip("#"))
    pPr.append(shd)

def set_cell_shading(cell, shading_hex):
    if not shading_hex:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shading_hex.lstrip("#"))
    tcPr.append(shd)

def set_alignment(paragraph, align_str):
    if not align_str:
        return
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = align_map.get(align_str.lower(), None)

def set_indents(paragraph, indents):
    if not indents:
        return
    pf = paragraph.paragraph_format
    if 'left' in indents:
        pf.left_indent = Pt(indents['left'])
    if 'right' in indents:
        pf.right_indent = Pt(indents['right'])
    if 'first_line' in indents:
        pf.first_line_indent = Pt(indents['first_line'])
    if 'hanging' in indents:
        pf.first_line_indent = Pt(-indents['hanging'])

def set_spacing(paragraph, spacing):
    if not spacing:
        return
    pf = paragraph.paragraph_format
    if 'space_before' in spacing:
        pf.space_before = Pt(spacing['space_before'])
    if 'space_after' in spacing:
        pf.space_after = Pt(spacing['space_after'])
    if 'line_spacing' in spacing:
        pf.line_spacing = spacing['line_spacing']

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    lv = len(hex_color)
    return tuple(int(hex_color[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))

def deep_copy_paragraph(source, target):
    target.style = source.style
    set_alignment(target, source.alignment.name.lower() if source.alignment else None)
    set_indents(target, get_paragraph_indents(source))
    set_spacing(target, get_paragraph_spacing(source))
    if shading := get_paragraph_shading(source):
        set_paragraph_shading(target, shading)
    target.clear()
    for run in source.runs:
        new_run = target.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        if color := get_font_color(run):
            r, g, b = hex_to_rgb(color)
            new_run.font.color.rgb = RGBColor(r, g, b)
        if shading := get_run_shading(run):
            set_run_shading(new_run, shading)

def set_table_borders(table):
    tbl = table._tbl
    borders_xml = (
        '<w:tblBorders {}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ).format(nsdecls('w'))
    tbl.tblPr.append(parse_xml(borders_xml))
    
def deep_copy_table(source_table, target_doc, color=None):
    target_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    set_table_borders(target_table)
    for r, source_row in enumerate(source_table.rows):
        for c, source_cell in enumerate(source_row.cells):
            target_cell = target_table.cell(r, c)
            if shading := get_cell_shading(source_cell):
                set_cell_shading(target_cell, '#8A084B')
            for source_p in source_cell.paragraphs:
                target_p = target_cell.add_paragraph()
                deep_copy_paragraph(source_p, target_p)
                if color:
                    for run in target_p.runs:
                        run.font.color.rgb = color
    return target_table

def extract_articles(doc):
    articles = {}
    blocks = list(iter_block_items(doc))
    current_article = None
    start = 0
    for i, block in enumerate(blocks):
        if isinstance(block, Paragraph) and ARTICLE_RE.match(block.text.strip()):
            if current_article:
                articles[current_article] = (start, i)
            current_article = block.text.strip()
            start = i
    if current_article:
        articles[current_article] = (start, len(blocks))
    return articles

def merge_gazette(old_doc, gov_doc, target_doc):
    for block in iter_block_items(old_doc):
        if isinstance(block, Paragraph) and "Sl. glasnik RS" in block.text:
            gov_line = next((q.text for q in iter_block_items(gov_doc) if isinstance(q, Paragraph) and "Sl. glasnik RS" in q.text), "")
            merged = merge_gazette_text(block.text, gov_line)
            for t_block in iter_block_items(target_doc):
                if isinstance(t_block, Paragraph) and "Sl. glasnik RS" in t_block.text:
                    deep_copy_paragraph(block, t_block)
                    t_block.text = merged
                    return True
    return False

def merge_gazette_text(old_text, new_text):
    m_old = GAZETTE_RE.search(old_text)
    m_new = GAZETTE_RE.search(new_text)
    if m_old and m_new:
        old_refs = set(ref.strip() for ref in m_old.group(1).split(' i '))
        new_refs = set(ref.strip() for ref in m_new.group(1).split(' i '))
        merged_refs = sorted(old_refs.union(new_refs), key=lambda x: (not x.isdigit(), x))
        return old_text.replace(m_old.group(0), f'(\"Sl. glasnik RS\", br. {" i ".join(merged_refs)})')
    return old_text

def apply_amendment_text(old_text, instruction):
    prompt = f"""
    You are an expert in Serbian legislative amendments. Apply the changes to the old article text exactly as per the instruction. If it says a stav 'prestaju da važe' (ceases to be valid), delete that stav from the text, append '*' to the article title. The 'stav' number refers to the (number)th paragraph after the title. Return only the updated text, using new lines for paragraphs/stavs, preserving structure.

    Example 1 (exact sample):
    Old text: 'Član 9 \n<p>Knjiženje poslovnih promena i događaja (u daljem tekstu: poslovnih promena) na računima imovine, obaveza, kapitala, prihoda i rashoda vrši se na osnovu verodostojnih računovodstvenih isprava. \nRačunovodstvena isprava predstavlja pisani dokument ili elektronski zapis o nastaloj poslovnoj promeni, koja obuhvata sve podatke potrebne za knjiženje u poslovnim knjigama tako da se iz računovodstvene isprave nedvosmisleno može saznati osnov, vrsta i sadržaj poslovne promene. \nFaktura (račun) kao računovodstvena isprava, u smislu ovog zakona, sastavlja se i dostavlja pravnim licima i preduzetnicima u elektronskom obliku i mora biti potvrđena od strane odgovornog lica koje svojim potpisom ili drugom identifikacionom oznakom (utvrđenom opštim aktom kojim pravno lice, odnosno preduzetnik uređuje organizaciju računovodstva) potvrđuje njenu verodostojnost. \nRačunovodstvena isprava sastavlja se u potrebnom broju primeraka na mestu i u vreme nastanka poslovne promene. \nRačunovodstvena isprava koja je sastavljena u jednom primerku može se otpremiti ako su podaci iz te isprave stalno dostupni. \nFotokopija računovodstvene isprave je osnov za knjiženje poslovne promene, pod uslovom da je na njoj navedeno mesto čuvanja originalne isprave i da je potvrđena od strane odgovornog lica koji svojim potpisom ili drugom identifikacionom oznakom potvrđuje njenu verodostojnost. \nRačunovodstvenom ispravom smatra se i isprava ispostavljena, odnosno primljena telekomunikacionim putem, kao i isprava ispostavljena, odnosno primljena putem servisa za elektronsku razmenu podataka (Electronic data Interchange - EDI). \nPošiljalac je odgovoran da podaci na ulazu u telekomunikacioni sistem budu zasnovani na računovodstvenim ispravama, kao i da čuva originalne računovodstvene isprave. \nKada se računovodstvena isprava prenosi putem servisa za elektronsku razmenu podataka, pružalac usluge elektronske razmene podataka dužan je da obezbedi integritet razmenjenih podataka.'
    Instruction: 'člana 9. stav 3. Zakona o računovodstvu prestaje da važe'
    Updated: 'Član 9*\nKnjiženje poslovnih promena i događaja (u daljem tekstu: poslovnih promena) na računima imovine, obaveza, kapitala, prihoda i rashoda vrši se na osnovu verodostojnih računovodstvenih isprava. \nRačunovodstvena isprava predstavlja pisani dokument ili elektronski zapis o nastaloj poslovnoj promeni, koja obuhvata sve podatke potrebne za knjiženje u poslovnim knjigama tako da se iz računovodstvene isprave nedvosmisleno može saznati osnov, vrsta i sadržaj poslovne promene. \nRačunovodstvena isprava sastavlja se u potrebnom broju primeraka na mestu i u vreme nastanka poslovne promene. \nRačunovodstvena isprava koja je sastavljena u jednom primerku može se otpremiti ako su podaci iz te isprave stalno dostupni. \nFotokopija računovodstvene isprave je osnov za knjiženje poslovne promene, pod uslovom da je na njoj navedeno mesto čuvanja originalne isprave i da je potvrđena od strane odgovornog lica koji svojim potpisom ili drugom identifikacionom oznakom potvrđuje njenu verodostojnost. \nRačunovodstvenom ispravom smatra se i isprava ispostavljena, odnosno primljena telekomunikacionim putem, kao i isprava ispostavljena, odnosno primljena putem servisa za elektronsku razmenu podataka (Electronic data Interchange - EDI). \nPošiljalac je odgovoran da podaci na ulazu u telekomunikacioni sistem budu zasnovani na računovodstvenim ispravama, kao i da čuva originalne računovodstvene isprave. \nKada se računovodstvena isprava prenosi putem servisa za elektronsku razmenu podataka, pružalac usluge elektronske razmene podataka dužan je da obezbedi integritet razmenjenih podataka.'

    Example 2 (exact sample):
    Old text: 'Član 64 \nOdredbe člana 4. stav 7, člana 32. stav 4. tačka 2), člana 39. stav 3. tačka 2) i člana 40. stav 5. tačka 4) ovog zakona primenjuju se od dana prijema Republike Srbije u Evropsku uniju. \nOdredbe člana 6. st. 13. i 14, člana 29, čl. 44-49, čl. 51. i 52. ovog zakona, počeće da se primenjuju od finansijskih izveštaja koji se sastavljaju na dan 31. decembra 2021. godine. \nOdredba člana 9. stav 3. ovog zakona primenjuje se počev od 1. januara 2022. godine.'
    Instruction: 'člana 64. stav 3. Zakona o računovodstvu prestaje da važe'
    Updated: 'Član 64*\nOdredbe člana 4. stav 7, člana 32. stav 4. tačka 2), člana 39. stav 3. tačka 2) i člana 40. stav 5. tačka 4) ovog zakona primenjuju se od dana prijema Republike Srbije u Evropsku uniju. \nOdredbe člana 6. st. 13. i 14, člana 29, čl. 44-49, čl. 51. i 52. ovog zakona, počeće da se primenjuju od finansijskih izveštaja koji se sastavljaju na dan 31. decembra 2021. godine.'

    Now apply:
    Old text:\n{old_text}\n\nInstruction:\n{instruction}
    """
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model="gpt-4.1-nano",
                messages=[{"role": "system", "content": "You are a precise legal document editor for Serbian laws."}, {"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=2000
            )
            content = response.choices[0].message.content.strip()
            lines = [line.strip() for line in content.splitlines() if line.strip()]
            if lines and len(lines) > 1 and '*' in lines[0]:
                logging.info("GPT success.")
                return lines
            raise ValueError("Invalid output")
        except Exception as e:
            logging.warning(f"GPT attempt {attempt+1}: {e}")
    logging.error("GPT failed. Falling back.")
    return old_text.splitlines()

def extract_amending_ref(gov_doc):
    article = ""
    law_name = ""
    gazette = ""
    for block in iter_block_items(gov_doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if ARTICLE_RE.match(text):
                article = text.upper()
            if "ZAKON O" in text.upper():
                law_name = text.upper()
            if "SL. GLASNIK RS" in text.upper():
                gazette = text.upper()
    if article and law_name and gazette:
        return f"[{article} {law_name} {gazette}]"
    return '[ČLAN 23 STAV 2 ZAKONA O ELEKTRONSKOM FAKTURISANJU ("SL. GLASNIK RS", BR. 44/2021)]'

def add_explanatory_table(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, '#CCCCCC')
    p = cell.add_paragraph()
    set_alignment(p, 'center')
    set_spacing(p, {'space_before': 5.0, 'space_after': 5.0, 'line_spacing': 1.0})
    texts = [
        ("Radi lakšeg sagledavanja izmena i dopuna propisa, nova sadržina odredaba data je ", True, None, None, False),
        ("zelenom", True, None, '#33FF33', False),
        (", prethodna ", True, None, None, False),
        ("crvenom", True, '#FF0000', None, False),
        (" bojom, a nepromenjene odredbe nisu posebno označene, tako da pregledanjem crno-zelenog teksta pregledate važeću, a crno-crvenog teksta, prethodnu verziju propisa. Prečišćen tekst bez crvenih i zelenih oznaka i dalje možete videti na tabu ", True, None, None, False),
        ("\"Tekst dokumenta\".", True, None, None, True)
    ]
    for text, bold, color, shading, italic in texts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Arial'
        if color:
            r, g, b = hex_to_rgb(color)
            run.font.color.rgb = RGBColor(r, g, b)
        if shading:
            set_run_shading(run, shading)

    # Second table (manual, no repeat)
    table2 = doc.add_table(rows=1, cols=1)
    cell2 = table2.cell(0, 0)
    set_cell_shading(cell2, '#000000')
    p = cell2.add_paragraph()
    set_alignment(p, 'center')
    set_spacing(p, {'line_spacing': 1.1})

    run = p.add_run("Propis - analitički prikaz promena")
    run.bold = True
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(255, 232, 191)

# GUI App class - only available when tkinter can be imported
try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext
    
    class App(tk.Tk):
        def __init__(self):
            super().__init__()
            self.title("Legal Document Processor")
            self.geometry("700x500")
            self.orig_path = self.amend_path = self.new_path = ""
            self.updated_doc = self.diff_doc = None

            frame = tk.Frame(self)
            frame.pack(pady=10, fill="x")

            tk.Label(frame, text="Original Law (.docx):").grid(row=0, column=0, sticky="w")
            self.entry_orig = tk.Entry(frame, width=60)
            self.entry_orig.grid(row=0, column=1)
            tk.Button(frame, text="Browse", command=self.select_original).grid(row=0, column=2)

            tk.Label(frame, text="Gov Changes (.docx):").grid(row=1, column=0, sticky="w")
            self.entry_amend = tk.Entry(frame, width=60)
            self.entry_amend.grid(row=1, column=1)
            tk.Button(frame, text="Browse", command=self.select_amendment).grid(row=1, column=2)

            tk.Label(frame, text="New Law (.docx) for Diff:").grid(row=2, column=0, sticky="w")
            self.entry_new = tk.Entry(frame, width=60)
            self.entry_new.grid(row=2, column=1)
            tk.Button(frame, text="Browse", command=self.select_new).grid(row=2, column=2)

            tk.Button(self, text="Process Part A: Generate New.docx", bg="#007ACC", fg="white", command=self.process_part_a).pack(pady=5, fill="x", padx=50)
            tk.Button(self, text="Process Part B: Generate Colored Diff.docx", bg="#28a745", fg="white", command=self.process_part_b).pack(pady=5, fill="x", padx=50)
            tk.Button(self, text="Save New.docx", command=lambda: self.save(self.updated_doc, "new")).pack(pady=5, fill="x", padx=50)
            tk.Button(self, text="Save Colored Diff.docx", command=lambda: self.save(self.diff_doc, "colored_diff")).pack(pady=5, fill="x", padx=50)

            self.log_text = scrolledtext.ScrolledText(self, height=10, state='disabled')
            self.log_text.pack(pady=10, fill="both", expand=True)
            self.update_log("Ready.")

        def update_log(self, msg):
            self.log_text.config(state='normal')
            self.log_text.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')} - {msg}\n")
            self.log_text.see(tk.END)
            self.log_text.config(state='disabled')

        def select_original(self):
            path = filedialog.askopenfilename(initialdir=DEFAULT_DIR, filetypes=[("Word Documents", "*.docx")])
            if path:
                self.orig_path = path
                self.entry_orig.delete(0, tk.END)
                self.entry_orig.insert(0, path)

        def select_amendment(self):
            path = filedialog.askopenfilename(initialdir=DEFAULT_DIR, filetypes=[("Word Documents", "*.docx")])
            if path:
                self.amend_path = path
                self.entry_amend.delete(0, tk.END)
                self.entry_amend.insert(0, path)

        def select_new(self):
            path = filedialog.askopenfilename(initialdir=DEFAULT_DIR, filetypes=[("Word Documents", "*.docx")])
            if path:
                self.new_path = path
                self.entry_new.delete(0, tk.END)
                self.entry_new.insert(0, path)

        def process_part_a(self):
            if not self.orig_path or not self.amend_path:
                messagebox.showerror("Error", "Select original and gov changes files first.")
                return
            self.update_log("Processing Part A...")
            logging.info("Starting Part A")

            orig = Document(self.orig_path)
            self.updated_doc = Document()
            for block in iter_block_items(orig):
                if isinstance(block, Paragraph):
                    p = self.updated_doc.add_paragraph()
                    deep_copy_paragraph(block, p)
                elif isinstance(block, Table):
                    deep_copy_table(block, self.updated_doc)

            merge_gazette(orig, Document(self.amend_path), self.updated_doc)

            articles = extract_articles(self.updated_doc)
            amend_doc = Document(self.amend_path)
            amendments = [p.text.strip() for p in amend_doc.paragraphs if "prestaju da važe" in p.text]

            for inst in amendments:
                matches = CHANGE_RE.findall(inst)
                for article_num, stav_num in matches:
                    aid = f"Član {article_num}"
                    if aid in articles:
                        s, e = articles[aid]
                        old_lines = '\n'.join(self.updated_doc.paragraphs[k].text for k in range(s, e))
                        new_lines = apply_amendment_text(old_lines, inst)
                        
                        # Remove old paras
                        for k in range(e - 1, s - 1, -1):
                            self.updated_doc.paragraphs[k]._element.getparent().remove(self.updated_doc.paragraphs[k]._element)
                        
                        # Insert new
                        for i, line in enumerate(new_lines):
                            new_p = self.updated_doc.add_paragraph(line)
                            original_format_idx = s + i
                            if original_format_idx < len(orig.paragraphs):
                                deep_copy_paragraph(orig.paragraphs[original_format_idx], new_p)
                            # Ensure article titles have center, Arial 12 bold
                            if ARTICLE_RE.match(line):
                                set_alignment(new_p, 'center')
                                for run in new_p.runs:
                                    run.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(12)

            self.update_log("Part A complete.")
            messagebox.showinfo("Success", "New.docx generated.")

        def process_part_b(self):
            if not self.orig_path or not self.new_path or not self.amend_path:
                messagebox.showerror("Error", "Select all files first.")
                return
            self.update_log("Processing Part B...")
            logging.info("Starting Part B")

            orig = Document(self.orig_path)
            new_d = Document(self.new_path)
            gov = Document(self.amend_path)
            self.diff_doc = Document()

            add_explanatory_table(self.diff_doc)

            # # Add adjusted title table once
            # for block in iter_block_items(new_d):
            #     if isinstance(block, Table):
            #         target_table = deep_copy_table(block, self.diff_doc)
            #         for row in target_table.rows:
            #             for cell in row.cells:
            #                 set_cell_shading(cell, '#8A084B')
            #         break

            # Prepare texts for diff, normalizing article titles by removing '*'
            orig_blocks = list(iter_block_items(orig))
            new_blocks = list(iter_block_items(new_d))
            orig_texts = []
            for b in orig_blocks:
                if isinstance(b, Paragraph):
                    text = b.text.strip()
                    orig_texts.append(text)
                elif isinstance(b, Table):
                    orig_texts.append('TABLE')
            new_texts = []
            for b in new_blocks:
                if isinstance(b, Paragraph):
                    text = b.text.strip()
                    text = text.replace('*', '')
                    new_texts.append(text)
                elif isinstance(b, Table):
                    new_texts.append('TABLE')

            matcher = SequenceMatcher(None, orig_texts, new_texts)

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == 'equal':
                    for k in range(i1, i2):
                        block = new_blocks[j1 + (k - i1)]
                        if isinstance(block, Paragraph):
                            p = self.diff_doc.add_paragraph()
                            deep_copy_paragraph(block, p)
                        elif isinstance(block, Table):
                            deep_copy_table(block, self.diff_doc)
                elif tag == 'delete':
                    for k in range(i1, i2):
                        block = orig_blocks[k]
                        if isinstance(block, Paragraph):
                            p = self.diff_doc.add_paragraph('[' + block.text + ']')
                            deep_copy_paragraph(block, p)
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                        elif isinstance(block, Table):
                            deep_copy_table(block, self.diff_doc, color=RGBColor(255, 0, 0))
                elif tag == 'insert':
                    for k in range(j1, j2):
                        block = new_blocks[k]
                        if isinstance(block, Paragraph):
                            p = self.diff_doc.add_paragraph('[' + block.text + ']')
                            deep_copy_paragraph(block, p)
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0, 204, 51)
                        elif isinstance(block, Table):
                            deep_copy_table(block, self.diff_doc, color=RGBColor(0, 204, 51))
                elif tag == 'replace':
                    for k in range(i1, i2):
                        block = orig_blocks[k]
                        if isinstance(block, Paragraph):
                            p = self.diff_doc.add_paragraph('[' + block.text + ']')
                            deep_copy_paragraph(block, p)
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                        elif isinstance(block, Table):
                            deep_copy_table(block, self.diff_doc, color=RGBColor(255, 0, 0))
                    for k in range(j1, j2):
                        block = new_blocks[k]
                        if isinstance(block, Paragraph):
                            p = self.diff_doc.add_paragraph('[' + block.text + ']')
                            deep_copy_paragraph(block, p)
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0, 204, 51)
                        elif isinstance(block, Table):
                            deep_copy_table(block, self.diff_doc, color=RGBColor(0, 204, 51))

            # Insert dynamic green reference before changed articles (collect positions first to avoid index shifts)
            diff_blocks = list(iter_block_items(self.diff_doc))
            insert_positions = []
            ref = extract_amending_ref(gov)
            for idx, block in enumerate(diff_blocks):
                if isinstance(block, Paragraph) and ARTICLE_RE.match(block.text.strip()) and '*' in block.text:
                    insert_positions.append(idx)
            for pos in reversed(insert_positions):
                ref_p = self.diff_doc.add_paragraph()
                set_alignment(ref_p, 'center')
                set_spacing(ref_p, {'space_before': 12.0, 'space_after': 6.0, 'line_spacing': 1.0})
                run = ref_p.add_run(ref)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 204, 51)
                self.diff_doc._body._element.insert(pos, ref_p._element)

            # Add "." spacer after Član 1 body
            diff_blocks = list(iter_block_items(self.diff_doc))  # Refresh list after insertions
            for idx, block in enumerate(diff_blocks):
                if isinstance(block, Paragraph) and block.text.strip() == "Član 1":
                    spacer = self.diff_doc.add_paragraph(".")
                    deep_copy_paragraph(block, spacer)
                    self.diff_doc._body._element.insert(idx + 2, spacer._element)
                    break

            self.update_log("Part B complete.")
            messagebox.showinfo("Success", "Colored diff generated.")

        def save(self, doc, kind):
            if not doc:
                messagebox.showerror("Error", f"Process {kind} first.")
                return
            out = filedialog.asksaveasfilename(initialdir=DEFAULT_DIR, defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if out:
                doc.save(out)
                self.update_log(f"{kind.capitalize()} saved to {out}")
                messagebox.showinfo("Saved", f"{kind.capitalize()} saved to {out}")

except ImportError:
    # tkinter not available (e.g., on Streamlit Cloud)
    print("GUI not available - tkinter not installed")

if __name__ == "__main__":
    try:
        App().mainloop()
    except NameError:
        print("GUI not available. Please use the Streamlit interface (st_ui.py) instead.")
